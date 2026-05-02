from typing import List, Dict, Any, Optional
from pathlib import Path
import re

try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table
except ImportError:
    raise ImportError("python-docx is required. Install it with: pip install python-docx")

from logger import get_logger
from config import get_config


class DocumentParser:
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        self.logger = get_logger(__name__)
        self._doc: Optional[Document] = None

    def parse(self) -> List[Dict[str, Any]]:
        self._validate_file()
        if self.file_path.suffix.lower() == '.txt':
            return self._parse_txt()
        else:
            self._load_document()
            return self._extract_content()

    def _validate_file(self) -> None:
        if not self.file_path.exists():
            raise DocumentParserError(f"File not found: {self.file_path}")
        if self.file_path.suffix.lower() not in ['.docx', '.txt']:
            raise DocumentParserError(f"Only .docx and .txt files are supported. Got: {self.file_path.suffix}")

    def _load_document(self) -> None:
        try:
            self._doc = Document(self.file_path)
            self.logger.info(f"Successfully loaded document: {self.file_path}")
        except Exception as e:
            raise DocumentParserError(f"Failed to load document: {e}")

    def _parse_txt(self) -> List[Dict[str, Any]]:
        content_blocks: List[Dict[str, Any]] = []
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            for i, line in enumerate(lines):
                line = line.strip()
                if line:
                    content_blocks.append({
                        'type': 'paragraph',
                        'text': line,
                        'style': 'Normal',
                        'line_number': i + 1
                    })
            
            self.logger.info(f"Successfully loaded TXT file: {self.file_path}")
            self.logger.info(f"Extracted {len(content_blocks)} lines")
            return content_blocks
        except Exception as e:
            raise DocumentParserError(f"Failed to parse TXT file: {e}")

    def _extract_content(self) -> List[Dict[str, Any]]:
        content_blocks: List[Dict[str, Any]] = []

        for element in self._doc.element.body:
            if element.tag.endswith('p'):
                para = self._find_paragraph_by_element(element)
                if para:
                    text = para.text.strip()
                    if text:
                        content_blocks.append({
                            'type': 'paragraph',
                            'text': text,
                            'style': para.style.name if para.style else 'Normal'
                        })

            elif element.tag.endswith('tbl'):
                table = self._find_table_by_element(element)
                if table:
                    table_data = self._extract_table(table)
                    if table_data:
                        content_blocks.append({
                            'type': 'table',
                            'data': table_data
                        })

        self.logger.info(f"Extracted {len(content_blocks)} content blocks")
        return content_blocks

    def _find_paragraph_by_element(self, element) -> Optional[Paragraph]:
        for para in self._doc.paragraphs:
            if para._element == element:
                return para
        return None

    def _find_table_by_element(self, element) -> Optional[Table]:
        for table in self._doc.tables:
            if table._element == element:
                return table
        return None

    def _extract_table(self, table: Table) -> List[List[str]]:
        data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                data.append(row_data)
        return data

    def extract_raw_text(self) -> str:
        self._validate_file()
        
        if self.file_path.suffix.lower() == '.txt':
            try:
                with open(self.file_path, 'r', encoding='utf-8') as f:
                    raw_text = f.read()
                self.logger.info(f"Extracted {len(raw_text)} characters from TXT file")
                return raw_text
            except Exception as e:
                raise DocumentParserError(f"Failed to extract text from TXT file: {e}")
        else:
            self._load_document()

            text_parts = []

            for para in self._doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            for table in self._doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            raw_text = '\n'.join(text_parts)
            self.logger.info(f"Extracted {len(raw_text)} characters of raw text")
            return raw_text

    def get_document_info(self) -> Dict[str, Any]:
        self._validate_file()
        
        if self.file_path.suffix.lower() == '.txt':
            info = {
                'file_name': self.file_path.name,
                'file_path': str(self.file_path.absolute()),
                'file_size': self.file_path.stat().st_size,
                'file_type': 'txt'
            }
            return info
        else:
            self._load_document()

            info = {
                'file_name': self.file_path.name,
                'file_path': str(self.file_path.absolute()),
                'file_size': self.file_path.stat().st_size,
                'paragraphs_count': len([p for p in self._doc.paragraphs if p.text.strip()]),
                'tables_count': len(self._doc.tables),
                'file_type': 'docx'
            }
            return info


class DocumentParserError(Exception):
    pass


def parse_document(file_path: str) -> List[Dict[str, Any]]:
    parser = DocumentParser(file_path)
    return parser.parse()


def extract_text_from_docx(file_path: str) -> str:
    parser = DocumentParser(file_path)
    return parser.extract_raw_text()